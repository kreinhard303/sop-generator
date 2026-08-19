"""Writes a generated SOP to a local .docx file, with inline screenshots.

Styled to Rednellac brand guidelines (REDNELLAC BRAND GUIDELINES REL 1.0):
corporate palette, Franklin Gothic / Arial system typography, and the
horizontal full-color logo in the header. Structural layout (title block
with byline/metadata table, divider rules, colored warning callouts, page
numbers) is modeled on Rednellac's own "Managed Services Project Process"
SOP, minus its off-brand font/color choices.
"""
from __future__ import annotations

import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from .sop_generator import Sop

# Corporate palette (brand guidelines p.22)
BLUE = RGBColor(0x00, 0x67, 0xB1)  # Rednellac Blue — dominant color
DARK_GRAY = RGBColor(0x71, 0x70, 0x73)  # Rednellac Dark Gray
LIGHT_GRAY = RGBColor(0xD8, 0xD9, 0xDA)  # Rednellac Light Gray
RED = RGBColor(0xC4, 0x12, 0x30)  # Rednellac Red — used sparingly, for warnings

# Corporate typeface is ITC Franklin Gothic Std; Word falls back to a system
# font wherever it isn't installed, so this is safe even unlicensed. Body
# text uses Arial per the guidelines' own "system typography" fallback rule
# (p.26): Arial for common applications whenever Franklin Gothic is unavailable.
FONT_HEADING = "ITC Franklin Gothic Std Demi"
FONT_BODY = "Arial"

LOGO_PATH = Path(__file__).resolve().parent.parent / "assets" / "rednellac_logo.png"
TAGLINE = "Building Salesforce for Builders."


def _slugify(title: str) -> str:
    slug = re.sub(r"[^a-zA-Z0-9]+", "-", title.strip().lower()).strip("-")
    return slug or "sop"


def _available_path(out_dir: Path, slug: str) -> Path:
    """Pick a filename that isn't locked by an already-open copy (e.g. the
    previous run's .docx still open in Word)."""
    candidate = out_dir / f"{slug}.docx"
    n = 2
    while candidate.exists():
        try:
            with open(candidate, "r+b"):
                break  # exists but not locked — fine to overwrite
        except PermissionError:
            candidate = out_dir / f"{slug}-{n}.docx"
            n += 1
    return candidate


def _set_font(font, *, name: str, size: Pt, color: RGBColor, bold: bool = False, italic: bool = False) -> None:
    font.name = name
    font.size = size
    font.color.rgb = color
    font.bold = bold
    font.italic = italic


def _add_bottom_border(paragraph, *, color: RGBColor, size: int = 8) -> None:
    """Add a horizontal rule under a paragraph (divider under headings)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "%02X%02X%02X" % (color[0], color[1], color[2]))
    pBdr.append(bottom)
    pPr.append(pBdr)


def _shade_cell(cell, hex_color: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _add_page_number_field(paragraph) -> None:
    """Insert a 'Page X of Y' field (Word computes it live)."""
    run = paragraph.add_run("Page ")
    _set_font(run.font, name=FONT_BODY, size=Pt(8), color=DARK_GRAY)

    def _field(instr: str):
        r = paragraph.add_run()
        _set_font(r.font, name=FONT_BODY, size=Pt(8), color=DARK_GRAY)
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        instr_el = OxmlElement("w:instrText")
        instr_el.set(qn("xml:space"), "preserve")
        instr_el.text = f" {instr} "
        sep = OxmlElement("w:fldChar")
        sep.set(qn("w:fldCharType"), "separate")
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        r._r.append(begin)
        r._r.append(instr_el)
        r._r.append(sep)
        r._r.append(end)

    _field("PAGE")
    mid = paragraph.add_run(" of ")
    _set_font(mid.font, name=FONT_BODY, size=Pt(8), color=DARK_GRAY)
    _field("NUMPAGES")


def _apply_base_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    _set_font(normal.font, name=FONT_BODY, size=Pt(11), color=RGBColor(0x33, 0x33, 0x33))

    h1 = doc.styles["Heading 1"]
    _set_font(h1.font, name=FONT_HEADING, size=Pt(28), color=BLUE, bold=True)
    h1.paragraph_format.space_after = Pt(2)

    h2 = doc.styles["Heading 2"]
    _set_font(h2.font, name=FONT_HEADING, size=Pt(14), color=BLUE, bold=True)
    h2.paragraph_format.space_before = Pt(20)
    h2.paragraph_format.space_after = Pt(8)

    for section in doc.sections:
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)


def _add_header(doc: Document) -> None:
    section = doc.sections[0]
    header_para = section.header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if LOGO_PATH.exists():
        header_para.add_run().add_picture(str(LOGO_PATH), width=Inches(1.6))


def _add_footer(doc: Document) -> None:
    section = doc.sections[0]
    footer_para = section.footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run(f"Rednellac  |  {TAGLINE}  |  ")
    _set_font(run.font, name=FONT_BODY, size=Pt(8), color=DARK_GRAY)
    _add_page_number_field(footer_para)


def _add_title_block(doc: Document, sop: Sop, *, source_url: str | None) -> None:
    title_p = doc.add_heading(sop.title, level=1)
    _add_bottom_border(title_p, color=BLUE, size=12)

    byline = doc.add_paragraph()
    byline.paragraph_format.space_before = Pt(6)
    byline.paragraph_format.space_after = Pt(14)
    today = date.today()
    run = byline.add_run(f"Generated {today:%B} {today.day}, {today:%Y}")
    _set_font(run.font, name=FONT_BODY, size=Pt(9.5), color=DARK_GRAY, italic=True)
    if source_url:
        run2 = byline.add_run(f"   ·   Source recording: {source_url}")
        _set_font(run2.font, name=FONT_BODY, size=Pt(9.5), color=DARK_GRAY, italic=True)

    if sop.purpose:
        purpose_p = doc.add_paragraph()
        purpose_p.paragraph_format.space_after = Pt(16)
        run = purpose_p.add_run(sop.purpose)
        _set_font(run.font, name=FONT_BODY, size=Pt(11.5), color=DARK_GRAY)


def create_sop_docx(sop: Sop, *, out_dir: Path, source_url: str | None = None) -> Path:
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = _available_path(out_dir, _slugify(sop.title))

    doc = Document()
    _apply_base_styles(doc)
    _add_header(doc)
    _add_footer(doc)
    _add_title_block(doc, sop, source_url=source_url)

    section = doc.sections[0]
    content_width = section.page_width - section.left_margin - section.right_margin

    if sop.prerequisites:
        h = doc.add_heading("Prerequisites", level=2)
        _add_bottom_border(h, color=LIGHT_GRAY, size=6)
        for item in sop.prerequisites:
            doc.add_paragraph(item, style="List Bullet")

    steps_h = doc.add_heading("Steps", level=2)
    _add_bottom_border(steps_h, color=LIGHT_GRAY, size=6)
    for step in sop.steps:
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(10)
        run = para.add_run(f"{step.number}. {step.instruction}")
        _set_font(run.font, name=FONT_BODY, size=Pt(11), color=RGBColor(0x33, 0x33, 0x33), bold=True)
        if step.notes:
            note = doc.add_paragraph(f"Note: {step.notes}")
            note.paragraph_format.left_indent = Inches(0.3)
            for run in note.runs:
                run.font.italic = True
                run.font.color.rgb = DARK_GRAY
        if step.frame_path:
            doc.add_picture(step.frame_path, width=content_width)

    if sop.warnings:
        h = doc.add_heading("Warnings", level=2)
        _add_bottom_border(h, color=RED, size=6)
        for run in h.runs:
            run.font.color.rgb = RED
        for item in sop.warnings:
            warn_p = doc.add_paragraph()
            warn_p.paragraph_format.space_after = Pt(6)
            run = warn_p.add_run(f"⚠  {item}")
            _set_font(run.font, name=FONT_BODY, size=Pt(11), color=RED, bold=True)

    doc.save(out_path)
    return out_path
